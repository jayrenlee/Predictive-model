"""
Word document exporter for YouTube Channel Engine session logs.
"""

from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_session_to_docx(
    session_log: list[dict], output_dir: str = "."
) -> str:
    """
    Write session_log to a .docx file.

    session_log entries: {"role": "user"|"assistant", "text": str}

    Returns the path to the saved file.
    """
    doc = Document()

    # ── Title ────────────────────────────────────────────────────────────────
    title = doc.add_heading("YouTube Channel Reverse Engineering", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    subtitle = doc.add_paragraph("Claude Channel Model v2 — Session Export")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(11)
    subtitle.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    timestamp = doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    timestamp.runs[0].font.size = Pt(9)
    timestamp.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    # ── Session content ───────────────────────────────────────────────────────
    state_pattern_map = {
        "STATE 1": "Channel to Clone",
        "STATE 2": "Branding Brief",
        "STATE 3": "Transcripts",
        "STATE 4": "Topic / Ideas",
        "STATE 5": "Style DNA Analysis",
        "STATE 6": "Script",
        "STATE 7": "Visual Style Profile",
        "STATE 8": "Image Prompts",
        "STATE 9": "Video Prompts",
        "STATE 10": "Thumbnail Analysis",
        "STATE 11": "Thumbnails",
        "STATE 12": "Export",
    }

    for entry in session_log:
        role = entry["role"]
        text = entry["text"]

        if role == "user":
            # User turns: small label + italic text
            p = doc.add_paragraph()
            run = p.add_run("You: ")
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x66, 0xCC)
            body = p.add_run(text)
            body.italic = True
            body.font.size = Pt(10)
            body.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        else:
            # Assistant turns: detect state headings inside the text
            lines = text.split("\n")
            for line in lines:
                stripped = line.strip()

                # Check for STATE headings
                is_state_heading = False
                for state_key, state_name in state_pattern_map.items():
                    if state_key in stripped.upper():
                        h = doc.add_heading(stripped, level=2)
                        h.runs[0].font.color.rgb = RGBColor(0x1A, 0x6B, 0x3C)
                        is_state_heading = True
                        break

                if is_state_heading:
                    continue

                # Section headings (lines ending with ":")
                if stripped.endswith(":") and len(stripped) < 60 and not stripped.startswith("●"):
                    h = doc.add_heading(stripped, level=3)
                    h.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                    continue

                # Bullet points
                if stripped.startswith(("●", "-", "•", "*")):
                    clean = stripped.lstrip("●-•* ").strip()
                    p = doc.add_paragraph(clean, style="List Bullet")
                    p.runs[0].font.size = Pt(10)
                    continue

                # Numbered list
                if stripped and stripped[0].isdigit() and "." in stripped[:3]:
                    p = doc.add_paragraph(stripped, style="List Number")
                    p.runs[0].font.size = Pt(10)
                    continue

                # Divider lines
                if set(stripped) <= set("─-=_"):
                    doc.add_paragraph()
                    continue

                # Blank line
                if not stripped:
                    doc.add_paragraph()
                    continue

                # Regular paragraph
                p = doc.add_paragraph(stripped)
                p.runs[0].font.size = Pt(10)

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph("— End of session export —")
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.size = Pt(9)
    footer_p.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # ── Save ──────────────────────────────────────────────────────────────────
    filename = f"youtube_engine_session_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    out_path = Path(output_dir) / filename
    doc.save(str(out_path))
    return str(out_path)
